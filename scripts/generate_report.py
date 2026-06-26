"""Generate a concise Arabic RTL DOCX report for the IR project — Damascus University.

Font: Tajawal (set for both Latin and complex-script so Word renders Arabic in it).
Kept intentionally concise.
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT = "Tajawal"


def set_rtl_paragraph(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>'))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_rtl_run(run):
    """Mark run RTL and force the Tajawal font for Latin + complex-script glyphs."""
    rPr = run._element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} val="1"/>'))
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)  # complex script = Arabic


def add_p(doc, text, bold=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT
    run.font.bold = bold
    set_rtl_run(run)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    set_rtl_paragraph(h)
    run = h.add_run(text)
    run.font.name = FONT
    set_rtl_run(run)
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 76, 153)
    else:
        run.font.size = Pt(12)
    return h


def add_table(doc, headers, rows, sizes=(10, 10)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")} val="1"/>'))

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_rtl_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT
        set_rtl_run(run)
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366" w:val="clear"/>'))
        run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            set_rtl_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = FONT
            set_rtl_run(run)
    doc.add_paragraph("")
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl_paragraph(p)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT
    set_rtl_run(run)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(2)
    return p


def generate_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Make the document default font Tajawal too.
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:cs"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)

    # ── Title page ────────────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph("")

    def centered(text, size, color, bold=True, font=FONT):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = font
        set_rtl_run(r)
        r.font.color.rgb = RGBColor(*color)
        return p

    centered("جامعة دمشق", 22, (0, 51, 102))
    centered("كلية الهندسة المعلوماتية", 16, (0, 76, 153))
    doc.add_paragraph("")
    centered("مشروع مقرر استرجاع المعلومات", 24, (0, 51, 102))
    doc.add_paragraph("")
    centered("محرك بحث لاسترجاع المعلومات", 18, (102, 0, 0))
    centered("IR Search Engine", 14, (80, 80, 80), bold=False)
    for _ in range(4):
        doc.add_paragraph("")
    centered("العام الدراسي 2025 - 2026", 12, (80, 80, 80), bold=False)
    doc.add_page_break()

    # ── 1. Intro + goal ──────────────────────────────────────────
    add_heading(doc, "1. مقدمة وهدف المشروع", level=1)
    add_p(doc,
        "استرجاع المعلومات هو علم البحث عن أكثر الوثائق صلة باستعلام المستخدم ضمن مجموعة ضخمة من "
        "النصوص. يهدف المشروع إلى بناء محرك بحث متكامل يعمل على مجموعتي بيانات كبيرتين، بأربعة نماذج "
        "استرجاع، مع بنية خدمية (SOA) وواجهة تفاعلية ونظام تقييم.")
    add_p(doc,
        "النماذج المدعومة: TF-IDF و BM25 والتمثيل الكثيف (Dense) والاسترجاع الهجين (Hybrid) بنوعيه "
        "التسلسلي والمتوازي، مع تحسين للاستعلام (تصحيح إملائي ومرادفات).")

    # ── 2. Datasets ──────────────────────────────────────────────
    add_heading(doc, "2. مجموعات البيانات", level=1)
    add_p(doc, "مجموعتان من مرجع BEIR، تُحمّلان تلقائياً عبر ir_datasets:")
    add_table(doc,
        ["الصلة", "استعلامات الاختبار", "عدد الوثائق", "المجال", "المجموعة"],
        [
            ["ثنائية", "10,000", "522,931", "أسئلة مكررة", "Quora"],
            ["متدرجة 0/1/2", "49", "382,545", "استرجاع الحجج", "Touché"],
        ])

    # ── 3. Architecture ──────────────────────────────────────────
    add_heading(doc, "3. البنية الخدمية (SOA)", level=1)
    add_p(doc,
        "يعمل النظام كمجموعة خدمات مستقلة (FastAPI) تتواصل عبر HTTP. البوابة (Gateway) على المنفذ "
        "8000 هي نقطة الدخول الوحيدة التي تتعامل معها الواجهة وتنسّق باقي الخدمات.")
    add_table(doc,
        ["الوظيفة", "المنفذ", "الخدمة"],
        [
            ["تنسيق خط البحث والتقييم", "8000", "البوابة"],
            ["تنظيف وتقطيع النصوص", "8001", "المعالجة المسبقة"],
            ["بناء الفهارس (Offline)", "8002", "الفهرسة"],
            ["البحث في الفهارس", "8003", "الاسترجاع"],
            ["الدمج وحساب المقاييس", "8004", "التصنيف والتقييم"],
            ["تصحيح إملائي ومرادفات", "8005", "تحسين الاستعلام"],
            ["الواجهة التفاعلية", "8501", "Streamlit"],
            ["مخزن الوثائق الأصلية", "27017", "MongoDB"],
        ])
    add_p(doc, "مسار الاستعلام:", bold=True)
    for s in [
        "تحسين الاستعلام (فقط في وضع basic+extra): تصحيح إملائي + مرادفات",
        "المعالجة المسبقة للاستعلام بنفس خط معالجة الوثائق",
        "الاسترجاع: إيجاد معرّفات أفضل الوثائق وفق النموذج المختار",
        "قراءة النص الأصلي لتلك المعرّفات من MongoDB وعرضه (لا يُعرض النص المجذّع)",
    ]:
        add_bullet(doc, s)

    # ── 4. Preprocessing ─────────────────────────────────────────
    add_heading(doc, "4. المعالجة المسبقة", level=1)
    add_p(doc,
        "تحوّل النص الخام إلى رموز نظيفة. تُطبَّق نفس الخطوات على الوثائق والاستعلامات لضمان التطابق:")
    for s in [
        "تطبيع Unicode وتحويل لأحرف صغيرة وإزالة الروابط والترقيم",
        "تقطيع إلى رموز (Tokenization) عبر NLTK",
        "إزالة كلمات التوقف (Stopwords)",
        "تجذيع (Stemming) عبر Snowball — أو تحليل صرفي عبر WordNet",
    ]:
        add_bullet(doc, s)

    # ── 5. Models ────────────────────────────────────────────────
    add_heading(doc, "5. نماذج الاسترجاع", level=1)
    add_table(doc,
        ["العيب", "الميزة", "الأساس", "النموذج"],
        [
            ["لا يفهم المرادفات", "بسيط وسريع", "تطابق كلمات + تشابه جيب التمام", "TF-IDF"],
            ["لا يفهم المعنى", "قوي، معاملات k1/b قابلة للضبط", "نموذج احتمالي محسّن", "BM25"],
            ["أبطأ، ذاكرة أكبر", "يفهم المعنى والمرادفات", "متجهات 384 بُعد + FAISS", "Dense"],
        ])
    add_p(doc,
        "تُبنى نماذج TF-IDF و BM25 من النص المعالَج، بينما يُبنى التمثيل الكثيف من النص الأصلي عبر "
        "النموذج all-MiniLM-L6-v2.")

    # ── 6. Hybrid + fusion ───────────────────────────────────────
    add_heading(doc, "6. الاسترجاع الهجين وطرق الدمج", level=1)
    add_p(doc,
        "التسلسلي: BM25 يجلب 100 مرشّح ثم Dense يعيد ترتيبهم. المتوازي: تشغيل BM25 و Dense معاً "
        "ثم دمج النتائج بإحدى الطرق:")
    for s in [
        "RRF: دمج بالترتيب العكسي score=Σ 1/(60+rank) — الأمتن ولا يحتاج معايرة",
        "Weighted: تطبيع الدرجات (min-max) وجمعها بأوزان قابلة للتعديل",
        "CombMNZ: مثل الموزون لكن يضرب بعدد النماذج التي أرجعت الوثيقة (يكافئ الاتفاق)",
    ]:
        add_bullet(doc, s)

    # ── 7. Query refinement ──────────────────────────────────────
    add_heading(doc, "7. تحسين الاستعلام", level=1)
    add_p(doc,
        "يعمل فقط في وضع basic+extra: تصحيح إملائي عبر pyspellchecker، وتوسيع بالمرادفات عبر "
        "WordNet. في وضع basic يُتخطّى هذا التحسين بالكامل.")

    # ── 8. Evaluation ────────────────────────────────────────────
    add_heading(doc, "8. التقييم", level=1)
    add_p(doc, "يُقاس الأداء على استعلامات الاختبار مقابل أحكام الصلة (qrels) عبر ir-measures:")
    add_table(doc,
        ["الوصف", "المقياس"],
        [
            ["جودة الترتيب الكلية", "MAP"],
            ["جودة أفضل 10 مع الصلة المتدرجة", "nDCG@10"],
            ["دقة أول 10 نتائج", "P@10"],
            ["نسبة الوثائق ذات الصلة ضمن 100", "Recall@100"],
        ])
    add_p(doc,
        "ملاحظة: تُقصر أحكام الصلة على الأسئلة المُختبَرة فقط، وإلا حُسبت الأسئلة غير المُشغَّلة "
        "كصفر فتنهار القيمة.", color=(153, 0, 0))

    # ── 9. Logging ───────────────────────────────────────────────
    add_heading(doc, "9. التسجيل والمراقبة", level=1)
    add_p(doc,
        "يسجّل عميل HTTP المشترك كل نداء بين الخدمات (الوجهة، ملخّص الحمولة، رمز الحالة، الزمن)، "
        "وتسجّل البوابة مراحل البحث الثلاث [1/3] تحسين [2/3] معالجة [3/3] استرجاع مع تحويل "
        "الاستعلام وعدد النتائج — مما يسهّل التشخيص (يظهر مثلاً تخطّي التحسين في وضع basic).")

    # ── 10. Technologies ─────────────────────────────────────────
    add_heading(doc, "10. التقنيات المستخدمة", level=1)
    add_table(doc,
        ["الاستخدام", "التقنية"],
        [
            ["لغة البرمجة", "Python 3.12"],
            ["خدمات API", "FastAPI"],
            ["مخزن الوثائق الأصلية", "MongoDB + pymongo"],
            ["الواجهة", "Streamlit"],
            ["البيانات", "ir_datasets"],
            ["TF-IDF / BM25", "scikit-learn / bm25s"],
            ["التمثيل الكثيف", "sentence-transformers + FAISS"],
            ["المقاييس", "ir-measures"],
            ["معالجة النصوص", "NLTK / WordNet"],
            ["تصحيح إملائي", "pyspellchecker"],
        ])

    # ── 11. Run ──────────────────────────────────────────────────
    add_heading(doc, "11. طريقة التشغيل", level=1)
    add_table(doc,
        ["الوظيفة", "الأمر"],
        [
            ["تثبيت المكتبات", "make install"],
            ["تحميل البيانات", "make download-data"],
            ["إدخال الوثائق إلى MongoDB", "make ingest"],
            ["بناء الفهارس (مرة واحدة)", "make build-index"],
            ["تشغيل كل الخدمات + الواجهة", "make run"],
            ["التقييم الكامل", "make eval"],
        ])
    add_p(doc, "ثم افتح المتصفح على: http://localhost:8501")

    # ── 12. Conclusion ───────────────────────────────────────────
    add_heading(doc, "12. الخلاصة", level=1)
    for s in [
        "أربعة نماذج استرجاع مع دمج هجين قابل للضبط",
        "بنية خدمية حقيقية عبر بوابة موحّدة، وتخزين النص الأصلي في MongoDB",
        "تحسين استعلام، تقييم بمقاييس معيارية، وتسجيل مفصّل للتشخيص",
        "تشغيل كامل النظام بأمر واحد: make run",
    ]:
        add_bullet(doc, s)

    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "reports", "تقرير_المشروع.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()
